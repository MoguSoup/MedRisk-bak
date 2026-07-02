from __future__ import annotations

import os
import sys
from pathlib import Path

deps = Path(os.environ.get("TEMP", "")) / "medrisk-docx-deps"
if deps.exists():
    sys.path.insert(0, str(deps))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO = Path(r"D:\CScode\code\aaa_production\MedRisk")
FINAL_DIR = REPO / "final_ret"
SRC_DOCX = FINAL_DIR / "唐俊_2026计科生产实习报告册.docx"
OUT_DOCX = FINAL_DIR / "唐俊_2026计科生产实习报告册_MedRisk填写版.docx"
SCREEN_DIR = FINAL_DIR / "screenshots"
NOTES_DIR = FINAL_DIR / "notes"
EMBED_DIR = FINAL_DIR / "render_check" / "embedded_images"


def block_text(element) -> str:
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def set_run_font(run, size: float, bold: bool = False, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def format_paragraph(paragraph, kind: str) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(24)
    if kind == "title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
    elif kind in {"h1", "h2", "h3"}:
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(8)
        pf.space_after = Pt(6)
    elif kind in {"caption", "reference"}:
        pf.first_line_indent = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(4 if kind == "reference" else 8)
        if kind == "caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraph(anchor, text: str = "", kind: str = "body"):
    paragraph = anchor.insert_paragraph_before("")
    format_paragraph(paragraph, kind)
    run = paragraph.add_run(text)
    if kind == "title":
        set_run_font(run, 16, True, "黑体")
    elif kind == "h1":
        set_run_font(run, 15, True, "黑体")
    elif kind == "h2":
        set_run_font(run, 14, True, "黑体")
    elif kind == "h3":
        set_run_font(run, 12, True, "黑体")
    elif kind == "caption":
        set_run_font(run, 10.5, False, "宋体")
    elif kind == "reference":
        set_run_font(run, 10.5, False, "宋体")
    else:
        set_run_font(run, 12, False, "宋体")
    return paragraph


def add_table_before(doc: Document, anchor, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or c_idx in {0, 2} else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, 10.5, r_idx == 0, "宋体")
    anchor._p.addprevious(table._tbl)
    add_paragraph(anchor, "", "body")


def add_figure(anchor, filename: str, caption: str) -> None:
    image_path = SCREEN_DIR / filename
    if not image_path.exists():
        return
    EMBED_DIR.mkdir(parents=True, exist_ok=True)
    embed_path = EMBED_DIR / (Path(filename).stem + ".jpg")
    if not embed_path.exists():
        from PIL import Image

        with Image.open(image_path) as img:
            rgb = img.convert("RGB")
            rgb.save(embed_path, "JPEG", quality=88, optimize=True)
    p = anchor.insert_paragraph_before("")
    format_paragraph(p, "caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(embed_path), width=Inches(5.75))
    add_paragraph(anchor, caption, "caption")


def remove_blocks_between(doc: Document, start_text: str, end_text: str) -> object:
    body = doc.element.body
    children = list(body)
    title_idx = max(i for i, child in enumerate(children) if "生产实习报告" in block_text(child))
    start_idx = next(i for i in range(title_idx + 1, len(children)) if start_text in block_text(children[i]))
    end_idx = next(i for i in range(start_idx + 1, len(children)) if end_text in block_text(children[i]))
    anchor_element = children[end_idx]
    for child in children[start_idx:end_idx]:
        body.remove(child)
    return anchor_element


def paragraph_from_element(doc: Document, element):
    for p in doc.paragraphs:
        if p._p is element:
            return p
    raise RuntimeError("anchor paragraph not found")


def main() -> None:
    if not SRC_DOCX.exists():
        raise FileNotFoundError(SRC_DOCX)
    doc = Document(str(SRC_DOCX))
    anchor_element = remove_blocks_between(doc, "一、实习任务与计划", "I职业素养评分表")
    anchor = paragraph_from_element(doc, anchor_element)

    add_paragraph(anchor, "一、实习任务与计划", "h1")
    add_table_before(
        doc,
        anchor,
        [
            ["序号", "任务计划", "时间"],
            ["1", "实习动员、项目背景调研，明确医疗健康风险预测平台的功能边界、教学演示定位和工程伦理要求。", "2025.6.15-2025.6.17"],
            ["2", "完成本地与云端开发环境搭建，梳理 Vue3、Spring Boot 3、FastAPI、Neo4j、Docker Compose 等技术栈。", "2025.6.18-2025.6.21"],
            ["3", "进行需求分析与系统设计，确定登录注册、疾病预测、知识库、图谱构建、智能问答、模型训练和审计管理等核心模块。", "2025.6.22-2025.6.25"],
            ["4", "完成核心功能编码与联调，重点修复启动、Neo4j、图谱批量构建、邮箱验证码、问答流式输出和审计 IP 等问题。", "2025.6.26-2025.7.1"],
            ["5", "完成测试验证和远端部署，执行后端、前端、模型服务、Compose 配置和云端健康检查，并截取系统运行结果。", "2025.7.2-2025.7.4"],
            ["6", "整理实习报告、结果截图、测试记录和项目总结，准备实习答辩材料。", "2025.7.5"],
        ],
    )

    add_paragraph(anchor, "二、实习过程", "h1")
    add_paragraph(anchor, "1 项目基础技术与方法", "h2")
    add_paragraph(anchor, "本次生产实习围绕 MedRisk 疾病风险预测与智能辅助诊断系统开展。项目采用前后端分离和多服务协同架构，将 Vue3 前端、Spring Boot 后端、FastAPI 模型服务、Neo4j 医学知识图谱和 Docker Compose 云端部署整合在同一平台中。通过该项目，我重点学习并实践了 Web 工程开发、医学知识图谱构建、机器学习模型训练、OpenAI-compatible 大模型接入、Linux 云服务器部署和软件测试验证等内容。")
    add_paragraph(anchor, "1.1 Spring Boot 与后端服务", "h3")
    add_paragraph(anchor, "后端使用 Spring Boot 3 组织认证授权、用户管理、疾病档案、病历案例、知识文档、知识图谱任务、审计日志、邮件验证码和报告生成等业务模块。实习中重点解决了本地 demo 环境启动失败的问题：由于 H2 数据库中已有旧 Flyway 迁移记录，迁移文件变化后产生 checksum 校验错误，导致后端 8080 端口未能正常启动，前端登录表现为 500。最终方案是只在本地 demo profile 中放宽 Flyway 校验，并保持 Linux 云端 MySQL 迁移策略严格，从而兼顾本地演示可用性和生产环境一致性。")
    add_paragraph(anchor, "1.2 Vue3 与前端交互", "h3")
    add_paragraph(anchor, "前端采用 Vue3、Vite 和 Element Plus，实现登录注册、疾病预测、智能问答、知识图谱、风险大屏、模型训练管理和系统审计等页面。界面优化过程中，重点处理了侧边栏滚动丢失、导航重复、文字过小、按钮换行、图表超宽、Logo 与文字重叠等问题，使管理员、医生和患者角色均能在清晰的功能入口中完成操作。")
    add_figure(anchor, "01-login-page.png", "图1 MedRisk 登录界面与品牌入口")
    add_paragraph(anchor, "1.3 FastAPI 模型服务与机器学习", "h3")
    add_paragraph(anchor, "模型服务使用 FastAPI 承载风险预测、训练任务、模型版本和模型启用功能。模型体系从单一 XGBoost 扩展为 Logistic Regression、RandomForest、ExtraTrees、HistGradientBoosting、XGBoost 等稳定模型，并预留 LightGBM、CatBoost、TabPFN、TabICL、FT-Transformer 等可选高级模型。训练任务保存模型类型、超参数、评估数据集和启用状态，模型服务重启后能够继续加载已启用版本。")
    add_paragraph(anchor, "1.4 Neo4j、GraphRAG 与大模型问答", "h3")
    add_paragraph(anchor, "Neo4j 用于存储医学知识图谱，为智能问答提供 GraphRAG 检索证据。问答模块接入 OpenAI-compatible 大模型接口，并支持阿里云百炼等服务；系统根据问题内容自动判断日常聊天、医学问答和图片输入，医学问答会检索知识图谱、疾病档案和病历案例，日常聊天则不检索知识库。前端采用流式输出和 Markdown 安全渲染，避免等待完整回答后才显示结果。")

    add_paragraph(anchor, "2 项目实践", "h2")
    add_paragraph(anchor, "主要阐述项目实践过程，重点说明本人参与的问题定位、方案设计、编码测试和部署验证工作。", "body")
    add_paragraph(anchor, "2.1 需求分析", "h2")
    add_paragraph(anchor, "MedRisk 需要解决的核心问题包括：第一，项目本地启动和云端部署不稳定，前端登录失败时难以定位真实原因；第二，Neo4j 连接状态只在图谱管理页刷新，切换页面后状态不一致；第三，知识图谱只能单文档构建，无法稳定完成全量与增量构建；第四，智能问答缺少流式输出、多模态输入和医学知识检索边界；第五，风险预测模型能力单一，缺少多模型训练、评估和启用闭环；第六，管理员界面存在功能重复、审计日志 IP 缺失和布局空白等问题。")
    add_paragraph(anchor, "2.1.1 系统目标", "h3")
    add_paragraph(anchor, "系统目标是构建一个可用于教学演示的医疗疾病风险预测与智能问答平台。功能上需要支持用户登录注册、邮箱验证码、疾病风险预测、报告生成、医学知识库管理、Neo4j 图谱构建与可视化、智能问答、多模态图片问答、模型训练管理、审计日志和云端部署。性能与可靠性上要求本地启动可诊断、远端 Docker Compose 可复现、核心接口健康可检查、单个模块失败不影响其他模块继续运行。")
    add_paragraph(anchor, "2.2 系统概要设计", "h2")
    add_paragraph(anchor, "系统采用前端、业务后端、模型服务和数据存储分层设计。前端通过 Nginx 或 Vite 访问 `/api`，由 Spring Boot 后端统一提供业务接口；模型训练和预测通过 FastAPI 模型服务完成；MySQL 保存用户、预测、训练、审计等结构化数据；Neo4j 保存医学知识实体、文档节点和关系；大模型通过 OpenAI-compatible 配置接入，避免将具体供应商写死在业务代码中。")
    add_figure(anchor, "02-admin-console.png", "图2 管理员控制台与系统概览")
    add_paragraph(anchor, "概要设计中还明确了本地和远端环境差异。本地 Windows 开发必须固定使用 conda MedRisk 环境，后端 demo profile 使用 H2 数据库；云端 med_tencent 使用 Docker Compose 运行 MySQL、Neo4j、后端、前端和模型服务。远端同步时不覆盖 `.env`、数据目录、上传目录、模型目录、日志目录和 Docker 卷，保证服务可持续运行。")
    add_paragraph(anchor, "2.3 系统详细设计", "h2")
    add_paragraph(anchor, "疾病预测模块采用表单化输入和模型解释输出。用户选择病种后填写年龄、血糖、BMI、血压、家族史等指标，后端调用模型服务返回风险等级、置信度和主要风险因素。结果页使用平面 SCI 风格可视化展示因素贡献，避免三维图在报告和页面中造成阅读负担。")
    add_figure(anchor, "04-disease-predict-result.png", "图3 疾病风险预测结果与因素解释")
    add_paragraph(anchor, "知识图谱模块分为图谱管理和图谱可视化。图谱管理负责 Neo4j 健康检查、全量重建、增量构建和失败原因展示；图谱可视化保留三维展示，提供节点类型、关系类型、布局、标签显示和可见范围筛选。构建逻辑采用逐文档处理，单个文档失败时记录原因并继续处理后续文档，避免批量任务整体中断。")
    add_figure(anchor, "07-knowledge-graph-3d.png", "图4 医学知识图谱三维可视化")
    add_paragraph(anchor, "智能问答模块使用会话列表和问答窗口布局，支持普通文字问题、医学问答、多模态图片输入和流式输出。系统会先判断问题类型，医学问题检索知识图谱、疾病档案和病历案例，日常聊天不检索。返回内容按标准 Markdown 渲染，推理过程和回答正文均可逐段显示，问诊内容可同步到导出的风险报告中。")
    add_figure(anchor, "06-qa-streaming-answer.png", "图5 智能问答流式输出与医学检索结果")
    add_paragraph(anchor, "模型训练管理模块包含模型版本、数据集管理、训练任务、模型评估、模型反馈和大模型配置。训练表单根据模型类型显示不同超参数，默认稳定模型可以直接训练，缺少高级依赖的模型会显示不可用原因。评估指标包括 Accuracy、Precision、Recall、Specificity、F1、ROC-AUC、PR-AUC、LogLoss、Brier Score 和混淆矩阵等，适合医学风险分类场景。")
    add_figure(anchor, "08-training-management.png", "图6 模型训练管理与任务配置")
    add_paragraph(anchor, "审计日志模块统一记录用户操作、资源类型、资源 ID、请求 IP 和时间。IP 解析优先使用 `X-Forwarded-For`，其次使用 `X-Real-IP` 和远端地址；对于流式问答等异步场景，显式传递请求 IP，避免后台线程丢失上下文。系统还预留 IP 黑名单，用于阻断异常来源访问。")
    add_figure(anchor, "09-audit-log-ip.png", "图7 审计日志中的请求 IP 记录")
    add_paragraph(anchor, "2.4 系统编码与测试", "h2")
    add_paragraph(anchor, "编码阶段按照模块逐步推进。启动脚本中固化 conda MedRisk 解析、端口清理、依赖检查和后端 demo profile；后端增加邮件验证码、LLM profile、请求 IP 上下文、图谱批量构建和报告同步等能力；前端重构侧栏、登录页、预测页、问答页和管理页；模型服务扩展多模型 adapter 和公开数据集评估能力。")
    add_paragraph(anchor, "测试阶段包括单元测试、构建测试、接口测试、远端健康检查和人工 UI 验证。后端使用 `mvn test` 验证认证、审计、图谱、问答和报告；前端使用 `npm run test` 与 `npm run build` 验证组件和生产构建；模型服务使用 conda MedRisk Python 执行 pytest；云端使用 `docker compose config` 和 `/api/health` 验证服务编排与后端健康状态。")
    add_figure(anchor, "10-health-check.png", "图8 云端后端健康检查结果")
    add_paragraph(anchor, "在远端验证时，`81.71.65.98:5241/api/health` 返回 UP，说明后端服务可用。由于截图时 `81.71.65.98:80` 返回 502，结果展示截图使用实际可用的 5241 入口完成，并在交付目录中保留入口检查记录。该问题说明云端反向代理仍需继续排查，但不影响本次平台主要功能展示。")
    add_paragraph(anchor, "2.5 结束语", "h2")
    add_paragraph(anchor, "2.5.1 项目工作小结", "h3")
    add_paragraph(anchor, "本次项目实践完成了 MedRisk 平台从启动修复、核心功能完善到云端验证的较完整迭代。项目不仅实现了疾病风险预测、智能问答、知识图谱、模型训练和审计日志等功能，还解决了本地启动失败、Neo4j 状态不一致、图谱无法批量构建、问答等待时间长、页面布局不稳定和审计 IP 缺失等实际工程问题。通过这些工作，系统的可用性、可维护性和演示效果均得到提升。")
    add_paragraph(anchor, "2.5.2 工程伦理审查与评价", "h3")
    add_paragraph(anchor, "医疗健康软件必须重视社会、健康、安全、隐私和法律因素。MedRisk 在页面和报告中持续提示系统仅用于教学演示和健康风险提示，不能替代医生诊断；用户密码、邮箱授权码和第三方 API Key 不写入源码和文档；审计日志记录请求 IP 便于追踪异常操作；大模型回答设置医疗边界，避免将普通模型输出误认为正式诊断。项目的积极意义在于帮助用户理解健康风险和医学知识，但其局限也很明确：模型预测依赖训练数据和特征质量，问答结果需要结合医生意见使用。")
    add_paragraph(anchor, "2.5.3 团队协作工作总结", "h3")
    add_paragraph(anchor, "在综合项目实践中，我主要承担问题排查、后端与模型服务联调、前端页面优化、云端部署验证和报告整理等工作。开发过程需要在需求、实现、测试和部署之间反复沟通，例如用户界面的侧栏、Logo、问答模式和可视化方案多次调整，最终形成更符合使用场景的版本。通过实践，我认识到团队项目中需求确认、版本记录、接口约定和测试闭环同样重要，不能只关注代码实现本身。")
    add_paragraph(anchor, "2.5.4 职业素养总结", "h3")
    add_paragraph(anchor, "本次实习使我进一步理解了计算机工程师在真实项目中的责任。面对启动失败、远端代理异常、图谱构建失败和问答接口不稳定等问题，需要先定位根因，再设计可复现的修复方案，并通过测试验证结果。对于涉及医疗和个人信息的系统，更要保持谨慎，避免泄露密钥和个人数据，避免夸大模型能力。整个过程提高了我分析复杂工程问题、编写文档、维护部署环境和进行系统性测试的能力。")
    add_paragraph(anchor, "2.6 参考文献", "h2")
    refs = [
        "[1] Craig Walls. Spring in Action[M]. Manning Publications, 2022.",
        "[2] Martin Fowler. Patterns of Enterprise Application Architecture[M]. Addison-Wesley, 2002.",
        "[3] Ian Goodfellow, Yoshua Bengio, Aaron Courville. Deep Learning[M]. MIT Press, 2016.",
        "[4] Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/.",
        "[5] Vue.js Guide[EB/OL]. https://vuejs.org/guide/.",
        "[6] FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/.",
        "[7] Neo4j Cypher Manual[EB/OL]. https://neo4j.com/docs/cypher-manual/.",
        "[8] Docker Compose Documentation[EB/OL]. https://docs.docker.com/compose/.",
        "[9] scikit-learn User Guide[EB/OL]. https://scikit-learn.org/stable/user_guide.html.",
        "[10] XGBoost Documentation[EB/OL]. https://xgboost.readthedocs.io/.",
        "[11] 阿里云百炼 OpenAI 兼容接口文档[EB/OL]. https://help.aliyun.com/zh/model-studio/.",
        "[12] CDC BRFSS Annual Survey Data[EB/OL]. https://www.cdc.gov/brfss/annual_data/.",
    ]
    for ref in refs:
        add_paragraph(anchor, ref, "reference")
    add_paragraph(anchor, "三、实习体会", "h1")
    add_paragraph(anchor, "通过本次生产实习，我对软件项目从需求分析到部署验证的完整过程有了更加真实的认识。课堂学习中，很多技术点是独立出现的，例如前端组件、后端接口、数据库设计、机器学习模型和 Linux 部署；而在 MedRisk 项目中，这些技术必须被组织成一个能够持续运行的系统。一个看似简单的登录失败，背后可能涉及前端代理、后端启动、数据库迁移、环境变量和日志定位等多个环节。因此，工程实践不能只看表面错误，而要通过日志、接口、配置和运行环境逐层排查。")
    add_paragraph(anchor, "在项目实践中，我印象最深的是本地和云端环境差异带来的问题。本地 Windows 需要固定 conda MedRisk 环境，后端 demo profile 使用 H2 数据库；云端 med_tencent 则通过 Docker Compose 运行 MySQL、Neo4j、后端、前端和模型服务。如果缺少清晰的启动规则，很容易出现同一项目在一台机器可运行、换到另一台机器就失败的情况。通过固化脚本、补充文档和增加健康检查，我认识到可复现性是工程项目质量的重要组成部分。")
    add_paragraph(anchor, "医疗健康场景也让我更加重视工程伦理。MedRisk 的风险预测和智能问答能够帮助用户理解疾病风险、学习医学知识，但它不能替代医生诊断。系统设计中必须明确免责声明，保护用户隐私，不泄露邮箱授权码、API Key 和会话 token，并对大模型输出设置边界。技术能力越强，越需要通过设计约束和审计机制降低误用风险。")
    add_paragraph(anchor, "这次实习还提升了我的测试意识。过去我更关注功能是否能在页面上操作成功，而这次项目要求同时验证后端单元测试、前端构建、模型服务 pytest、Docker Compose 配置、远端健康接口和 UI 截图结果。只有把这些验证连起来，才能说明功能真正可交付。特别是智能问答、图谱构建和模型训练这类跨服务功能，如果只验证单个页面，很难发现异步线程、请求 IP、Neo4j 连接或外部模型超时等隐藏问题。")
    add_paragraph(anchor, "总体来说，本次生产实习让我从“写功能”转向“交付系统”的视角。MedRisk 项目涉及前端交互、后端业务、模型服务、图数据库、云端部署和文档报告等多个方面，我在解决问题的过程中学习了如何拆解需求、定位根因、设计方案、记录版本变化并用测试证明结果。后续如果继续完善该项目，我希望进一步优化多疾病统一预测模型、补充更多公开医学数据集评估、完善 IP 黑名单和大模型配置管理，并增加端到端自动化测试，使系统更加稳定、可信和易维护。")

    doc.save(str(OUT_DOCX))
    summary = f"""# 生产实习报告生成记录

- 模板源文件：{SRC_DOCX}
- 输出 DOCX：{OUT_DOCX}
- 正文主题：MedRisk 疾病风险预测与智能问答平台
- 插入截图目录：{SCREEN_DIR}
- 说明：评分表、成绩表、企业评价表保持模板原样；未写入任何真实密钥或授权码。
"""
    (NOTES_DIR / "report-material-summary.md").write_text(summary, encoding="utf-8")


if __name__ == "__main__":
    main()
