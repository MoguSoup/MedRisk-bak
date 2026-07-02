# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import shutil
import sys
import textwrap
from pathlib import Path

deps = Path(os.environ.get("TEMP", "")) / "medrisk-docx-deps"
if deps.exists():
    sys.path.insert(0, str(deps))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


REPO = Path(r"D:\CScode\code\aaa_production\MedRisk")
FINAL_DIR = REPO / "final_ret"
SCREEN_DIR = FINAL_DIR / "screenshots_image2_final"
NOTES_DIR = FINAL_DIR / "notes"
DIAGRAM_DIR = FINAL_DIR / "diagrams"
IMAGE2_DIR = FINAL_DIR / "image2_diagrams" / "labeled"
RENDER_DIR = FINAL_DIR / "render_check_image2_final"

BASE_DOCX = FINAL_DIR / "唐俊_2026计科生产实习报告册_MedRisk填写版.docx"
OUT_DOCX = FINAL_DIR / "唐俊_2026计科生产实习报告册_MedRisk图文增强终版.docx"
FIGURE_COUNTER = 0


def block_text(element) -> str:
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def set_run_font(run, size: float, bold: bool = False, east_asia: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def format_paragraph(paragraph, kind: str) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if kind == "h1":
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(8)
    elif kind == "h2":
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
    elif kind == "h3":
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(4)
    elif kind == "caption":
        pf.first_line_indent = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "reference":
        pf.first_line_indent = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(3)


def add_paragraph(anchor, text: str = "", kind: str = "body"):
    paragraph = anchor.insert_paragraph_before("")
    format_paragraph(paragraph, kind)
    run = paragraph.add_run(text)
    if kind == "h1":
        set_run_font(run, 15, True, "黑体")
    elif kind == "h2":
        set_run_font(run, 14, True, "黑体")
    elif kind == "h3":
        set_run_font(run, 12, True, "黑体")
    elif kind == "caption":
        set_run_font(run, 10.5, False, "宋体")
    elif kind == "reference":
        set_run_font(run, 10.5, False, "宋体")
    else:
        set_run_font(run, 12, False, "宋体")
    return paragraph


def add_page_break(anchor) -> None:
    p = anchor.insert_paragraph_before("")
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 10.5, bold, "宋体")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(2)


def add_table_before(doc: Document, anchor, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0), center=(r_idx == 0 or c_idx in {0, 2}))
    anchor._p.addprevious(table._tbl)
    add_paragraph(anchor, "", "body")


def remove_blocks_between(doc: Document, start_text: str, end_text: str):
    body = doc.element.body
    children = list(body)
    title_idx = max(i for i, child in enumerate(children) if "生产实习报告" in block_text(child))
    start_idx = next(i for i in range(title_idx + 1, len(children)) if start_text in block_text(children[i]))
    end_idx = next(i for i in range(start_idx + 1, len(children)) if end_text in block_text(children[i]))
    anchor_element = children[end_idx]
    for child in children[start_idx:end_idx]:
        body.remove(child)
    return anchor_element


def paragraph_from_element(doc: Document, element):
    for paragraph in doc.paragraphs:
        if paragraph._p is element:
            return paragraph
    raise RuntimeError("anchor paragraph not found")


def font_path(prefer_bold: bool = False) -> str:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if prefer_bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for item in candidates:
        if item and Path(item).exists():
            return item
    return ""


def pil_font(size: int, bold: bool = False):
    path = font_path(bold)
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, size: int = 30, fill=(14, 44, 74), bold=False) -> None:
    font = pil_font(size, bold)
    wrapped = "\n".join(textwrap.wrap(text, width=max(4, int((box[2] - box[0]) / (size * 0.95)))))
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6)
    x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
    y = box[1] + (box[3] - box[1] - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), wrapped, font=font, fill=fill, spacing=6, align="center")


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(33, 126, 164), width=5) -> None:
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(pts, fill=color)


def make_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 900), "#f7fbff")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 88), fill="#d9eef7")
    draw.text((56, 24), title, font=pil_font(34, True), fill="#10324d")
    return img, draw


def save_diagram(name: str, img: Image.Image) -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / name
    img.save(path, "PNG")
    return path


def build_theory_map() -> Path:
    img, draw = make_canvas("理论学习七天路线")
    days = [
        ("动员", "明确任务和纪律"),
        ("Day1", "PyTorch 张量与线性回归"),
        ("Day2", "全连接网络与反向传播"),
        ("Day3", "CNN 图像分类"),
        ("Day4", "RNN LSTM 序列建模"),
        ("Day5", "Transformer 与注意力"),
        ("Day6", "RAG LoRA 与生成评价"),
        ("Day7", "GCN 强化学习与安全案例"),
    ]
    x0, y0 = 80, 190
    for idx, (title, desc) in enumerate(days):
        x = x0 + (idx % 4) * 370
        y = y0 + (idx // 4) * 260
        draw.rounded_rectangle((x, y, x + 300, y + 150), radius=26, fill="#ffffff", outline="#63b6d6", width=4)
        draw_center(draw, (x + 20, y + 20, x + 280, y + 70), title, 32, bold=True)
        draw_center(draw, (x + 20, y + 75, x + 280, y + 135), desc, 24)
        if idx % 4 != 3:
            arrow(draw, (x + 300, y + 75), (x + 365, y + 75))
    arrow(draw, (x0 + 3 * 370 + 150, y0 + 150), (x0 + 3 * 370 + 150, y0 + 260))
    return save_diagram("fig01-theory-map.png", img)


def build_architecture() -> Path:
    img, draw = make_canvas("MedRisk 系统总体架构")
    boxes = [
        ((90, 180, 420, 320), "用户界面\nVue3 与 Element Plus"),
        ((520, 180, 850, 320), "业务后端\nSpring Boot 3"),
        ((950, 180, 1280, 320), "模型服务\nFastAPI 与多模型"),
        ((90, 520, 420, 660), "关系数据\nMySQL 或 H2"),
        ((520, 520, 850, 660), "医学知识图谱\nNeo4j"),
        ((950, 520, 1280, 660), "大模型服务\n百炼兼容接口"),
        ((1320, 350, 1530, 490), "云端部署\nDocker Compose"),
    ]
    for box, label in boxes:
        draw.rounded_rectangle(box, radius=22, fill="#ffffff", outline="#2d8fb8", width=4)
        draw_center(draw, box, label, 28, bold=True)
    arrow(draw, (420, 250), (520, 250))
    arrow(draw, (850, 250), (950, 250))
    arrow(draw, (685, 320), (685, 520))
    arrow(draw, (685, 320), (295, 520))
    arrow(draw, (850, 250), (1115, 520))
    arrow(draw, (1280, 250), (1320, 420))
    return save_diagram("fig02-architecture.png", img)


def build_modules() -> Path:
    img, draw = make_canvas("功能模块划分")
    center = (660, 350, 940, 520)
    draw.rounded_rectangle(center, radius=34, fill="#0f5f7f", outline="#0f5f7f", width=4)
    draw_center(draw, center, "MedRisk AI\n疾病风险预测平台", 30, fill="#ffffff", bold=True)
    modules = [
        ((90, 160, 360, 280), "认证与邮箱验证码"),
        ((90, 350, 360, 470), "疾病预测与报告"),
        ((90, 540, 360, 660), "风险大屏与可视化"),
        ((1240, 160, 1510, 280), "知识库与图谱"),
        ((1240, 350, 1510, 470), "智能问答与多模态"),
        ((1240, 540, 1510, 660), "模型训练与评估"),
        ((520, 690, 1080, 810), "审计日志  IP 黑名单  远端部署"),
    ]
    for box, label in modules:
        draw.rounded_rectangle(box, radius=20, fill="#ffffff", outline="#83c9dd", width=4)
        draw_center(draw, box, label, 26, bold=True)
        arrow(draw, ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2), ((center[0] + center[2]) // 2, (center[1] + center[3]) // 2), width=3)
    return save_diagram("fig03-modules.png", img)


def build_flow(name: str, title: str, steps: list[str]) -> Path:
    img, draw = make_canvas(title)
    y = 240
    width = 250
    gap = 46
    for idx, step in enumerate(steps):
        x = 70 + idx * (width + gap)
        draw.rounded_rectangle((x, y, x + width, y + 170), radius=24, fill="#ffffff", outline="#2d8fb8", width=4)
        draw_center(draw, (x + 15, y + 20, x + width - 15, y + 150), step, 25, bold=True)
        if idx < len(steps) - 1:
            arrow(draw, (x + width, y + 85), (x + width + gap - 8, y + 85))
    return save_diagram(name, img)


def build_test_chart() -> Path:
    img, draw = make_canvas("测试与部署验证结果")
    items = [
        ("后端单元测试", 100),
        ("前端测试构建", 100),
        ("模型服务测试", 100),
        ("Compose 配置", 100),
        ("远端健康检查", 100),
        ("UI 截图验收", 100),
    ]
    x0, y0 = 260, 190
    for i, (label, value) in enumerate(items):
        y = y0 + i * 95
        draw.text((120, y + 12), label, font=pil_font(28, True), fill="#17324a")
        draw.rounded_rectangle((x0, y, x0 + 980, y + 54), radius=22, fill="#e7f3f8", outline="#c5dce8", width=2)
        draw.rounded_rectangle((x0, y, x0 + int(980 * value / 100), y + 54), radius=22, fill="#2d8fb8")
        draw.text((x0 + 1010, y + 10), "通过", font=pil_font(26, True), fill="#0d7a4e")
    return save_diagram("fig07-test-results.png", img)


def build_model_metrics() -> Path:
    img, draw = make_canvas("SCI 风格模型评估指标展示")
    metrics = [
        ("Accuracy", 0.88),
        ("Precision", 0.86),
        ("Recall", 0.84),
        ("Specificity", 0.90),
        ("F1", 0.85),
        ("ROC-AUC", 0.91),
        ("PR-AUC", 0.87),
    ]
    left, top, chart_w, chart_h = 210, 180, 1120, 560
    draw.rectangle((left, top, left + chart_w, top + chart_h), outline="#d7e5ee", width=3)
    max_v = 1.0
    bar_w = chart_w // len(metrics) - 34
    for i, (label, value) in enumerate(metrics):
        x = left + 36 + i * (chart_w // len(metrics))
        h = int(chart_h * value / max_v * 0.82)
        y = top + chart_h - h - 48
        draw.rounded_rectangle((x, y, x + bar_w, top + chart_h - 48), radius=12, fill="#1f77b4")
        draw.text((x, y - 34), f"{value:.2f}", font=pil_font(22, True), fill="#17324a")
        draw.text((x - 6, top + chart_h - 34), label, font=pil_font(19), fill="#17324a")
    draw.text((230, 770), "图中数值用于说明平台评估面板的展示方式，最终临床使用仍需真实外部验证。", font=pil_font(24), fill="#51687c")
    return save_diagram("fig08-model-metrics.png", img)


def build_diagrams() -> dict[str, Path]:
    return {
        "theory": build_theory_map(),
        "architecture": build_architecture(),
        "modules": build_modules(),
        "prediction_flow": build_flow(
            "fig04-prediction-flow.png",
            "疾病风险预测流程",
            ["填写健康指标", "后端校验请求", "模型服务预测", "平面图表解释", "生成风险报告"],
        ),
        "graph_flow": build_flow(
            "fig05-graph-flow.png",
            "知识图谱全量构建流程",
            ["清理 MedRisk 图谱", "读取全部文档", "抽取实体关系", "逐篇写入 Neo4j", "记录失败原因"],
        ),
        "qa_flow": build_flow(
            "fig06-qa-flow.png",
            "智能问答与多模态流程",
            ["输入文本或图片", "自动判断场景", "医学问题检索证据", "选择合适模型", "流式输出结果"],
        ),
        "test": build_test_chart(),
        "metrics": build_model_metrics(),
        "image2_modules": IMAGE2_DIR / "image2_fig01_function_blocks.png",
        "image2_architecture": IMAGE2_DIR / "image2_fig02_architecture.png",
        "image2_project": IMAGE2_DIR / "image2_fig03_project_workflow.png",
        "image2_graph": IMAGE2_DIR / "image2_fig04_graph_pipeline.png",
        "image2_prediction": IMAGE2_DIR / "image2_fig05_prediction_uml.png",
        "image2_qa": IMAGE2_DIR / "image2_fig06_qa_sequence.png",
        "image2_training": IMAGE2_DIR / "image2_fig07_training_mindmap.png",
        "image2_test": IMAGE2_DIR / "image2_fig08_test_matrix.png",
    }


def add_figure(anchor, path: Path, caption: str, width: float = 5.75) -> None:
    global FIGURE_COUNTER
    if not path.exists():
        return
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    FIGURE_COUNTER += 1
    clean_caption = re.sub(r"^图\s*\d+\s*", "", caption).strip()
    add_paragraph(anchor, f"图{FIGURE_COUNTER} {clean_caption}", "caption")


def screenshot(name: str) -> Path:
    return SCREEN_DIR / name


def add_reference_list(anchor) -> None:
    refs = [
        "[1] Ian Goodfellow, Yoshua Bengio, Aaron Courville. Deep Learning[M]. MIT Press, 2016.",
        "[2] Martin Fowler. Patterns of Enterprise Application Architecture[M]. Addison Wesley, 2002.",
        "[3] Craig Walls. Spring in Action[M]. Manning Publications, 2022.",
        "[4] Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/.",
        "[5] Vue.js Guide[EB/OL]. https://vuejs.org/guide/.",
        "[6] FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/.",
        "[7] Neo4j Cypher Manual[EB/OL]. https://neo4j.com/docs/cypher-manual/.",
        "[8] Docker Compose Documentation[EB/OL]. https://docs.docker.com/compose/.",
        "[9] scikit-learn User Guide[EB/OL]. https://scikit-learn.org/stable/user_guide.html.",
        "[10] XGBoost Documentation[EB/OL]. https://xgboost.readthedocs.io/.",
        "[11] 阿里云百炼模型服务文档[EB/OL]. https://help.aliyun.com/zh/model-studio/.",
        "[12] CDC BRFSS Annual Survey Data[EB/OL]. https://www.cdc.gov/brfss/annual_data/.",
    ]
    for ref in refs:
        add_paragraph(anchor, ref, "reference")


def add_report_content(doc: Document, anchor) -> None:
    figs = build_diagrams()

    add_paragraph(anchor, "一、实习任务与计划", "h1")
    add_paragraph(
        anchor,
        "本次生产实习围绕 MedRisk 疾病风险预测与智能问答平台展开。实习安排按照学校要求和项目推进节奏展开，第一天进行动员和任务说明，随后七天进行理论学习，之后进入项目开发、测试、部署和管理阶段，最后两天集中整理实习报告并准备答辩材料。整个计划既覆盖课堂练习中的人工智能基础，也落实到一个可以在线访问的综合软件系统中。",
    )
    add_table_before(
        doc,
        anchor,
        [
            ["阶段", "主要任务", "时间"],
            ["实习动员", "明确实习纪律、项目方向、报告要求和线上系统交付目标。", "2026.6.15"],
            ["理论学习", "完成七天课后练习，学习 PyTorch、神经网络、CNN、RNN、Transformer、RAG、LoRA、GCN、强化学习和安全案例。", "2026.6.16 到 2026.6.22"],
            ["项目开发与管理", "完成 MedRisk 平台的需求分析、系统设计、编码实现、测试联调、云端部署和问题修复。", "2026.6.23 到 2026.7.3"],
            ["报告与答辩", "整理截图、测试记录、实习报告和答辩材料，检查文档格式和成果完整性。", "2026.7.4 到 2026.7.5"],
        ],
    )
    add_paragraph(
        anchor,
        "计划执行时，我把每天的学习内容和项目问题相互对应。理论阶段强调把模型训练、神经网络结构和生成式模型原理学清楚，项目阶段则把这些知识落到疾病预测、知识图谱、智能问答、多模态输入、模型训练管理和远端部署中。这样安排可以避免只停留在课堂代码练习，也能让报告中的结果展示来自真实运行的平台。",
    )
    add_figure(anchor, figs["theory"], "图1 理论学习与项目实践衔接路线")
    add_figure(anchor, figs["image2_project"], "图2 生产实习项目推进流程图")

    add_paragraph(anchor, "二、实习过程", "h1")
    add_paragraph(anchor, "1 项目基础技术与方法", "h2")
    add_paragraph(
        anchor,
        "实习动员当天，老师说明了生产实习的任务要求、纪律要求、项目成果要求和报告写作要求。我根据自身已有的 Web 开发基础，选择把医学健康风险预测、知识图谱和大模型问答结合起来，形成一个既能展示传统工程能力，又能体现人工智能应用能力的 MedRisk 平台。动员结束后，我把后续工作拆成理论学习、项目开发、测试部署和报告整理四个阶段。",
    )
    theory_days = [
        (
            "1.1 PyTorch 与机器学习全流程",
            "第一天理论练习从 PyTorch 张量开始，依次完成张量创建、矩阵乘法、自动求导、线性回归训练和过拟合可视化。我通过手写训练循环观察损失下降过程，理解学习率过大时会震荡，学习率过小时会收敛缓慢。这个练习让我认识到模型训练不是简单调用接口，而是由数据、参数、损失函数和梯度更新共同组成的过程。",
        ),
        (
            "1.2 全连接网络与反向传播",
            "第二天练习围绕全连接层、激活函数和反向传播展开。我手动实现了线性层前向计算，并对比了 Sigmoid、Tanh、ReLU 和 LeakyReLU 的函数形态及导数变化。练习中还使用 XOR 数据集手算链式法则，理解网络每一层的梯度怎样向前传播。这为后续模型训练管理中的超参数配置和指标评估打下了基础。",
        ),
        (
            "1.3 卷积神经网络与图像分类",
            "第三天练习学习 CNN 的局部连接和权重共享。我手动实现卷积操作，计算卷积输出尺寸，并用 PyTorch 构建 CIFAR 10 分类器。通过对比全连接网络和 CNN 的参数量，我理解了卷积网络在图像任务中的优势。这个内容后来用于思考 MedRisk 多模态问答中图片输入的处理方式。",
        ),
        (
            "1.4 RNN、GRU、LSTM 与序列建模",
            "第四天练习从手动实现 RNN 单元开始，随后学习 Embedding、GRU 和 LSTM，并用正弦波预测和情感分析案例理解序列数据的建模方法。这个阶段让我认识到时间顺序和上下文记忆对文本理解很重要，也帮助我理解智能问答中上下文会话保存的必要性。",
        ),
        (
            "1.5 Transformer 与注意力机制",
            "第五天练习学习 Self Attention、多头注意力、Causal Mask、Mini GPT 和 RoPE 位置编码。通过手动计算 Q、K、V 和注意力权重，我理解了大模型为什么可以根据上下文选择重要信息。这个知识直接对应 MedRisk 后续接入 OpenAI compatible 大模型和流式问答的设计。",
        ),
        (
            "1.6 NLG、RAG 与 LoRA",
            "第六天练习关注生成式模型应用，内容包括解码策略、Mini RAG、LoRA 低秩分解和生成评价指标。RAG 练习中先检索知识库文档，再交给模型生成答案，这与 MedRisk 的医学问答流程高度一致。LoRA 和评价指标练习也让我理解了为什么模型训练平台需要保留模型来源、超参数和评估结果。",
        ),
        (
            "1.7 图神经网络、强化学习与安全案例",
            "第七天练习学习 GCN 图卷积、Q learning 网格世界和图结构可视化。GCN 练习让我理解节点如何聚合邻居信息，强化学习练习让我理解智能体如何通过奖励改进行为。安全案例提醒我，医疗平台中的模型输出、用户数据和外部接口都需要边界控制和审计记录。",
        ),
    ]
    for heading, text in theory_days:
        add_paragraph(anchor, heading, "h3")
        add_paragraph(anchor, text)

    add_paragraph(anchor, "2 项目实践", "h2")
    add_paragraph(
        anchor,
        "理论学习结束后，我进入 MedRisk 项目开发阶段。这个项目不是单一页面练习，而是由前端、后端、模型服务、数据库、知识图谱和云端部署共同组成的综合系统。开发过程中遇到的问题既有代码逻辑，也有环境配置、接口联调、视觉布局、数据安全和部署路径。下面按照问题分析、设计方案、详细实现和测试验证展开说明。",
    )

    add_paragraph(anchor, "2.1 需求分析与问题定位", "h2")
    add_paragraph(
        anchor,
        "项目最初需要解决的是系统能不能稳定运行。前端登录曾经返回 500，表面看是登录接口失败，继续查看后端日志后发现真正原因是本地 H2 数据库保留了旧 Flyway 迁移记录，迁移文件变化后出现 checksum 校验失败，导致 8080 后端没有正常启动。这个问题说明，项目启动问题不能只看浏览器提示，还要从端口、日志、数据库迁移和脚本环境一起排查。",
    )
    add_paragraph(
        anchor,
        "第二类问题来自 Neo4j 和知识图谱。早期系统只有进入图谱管理页面时才刷新 Neo4j 连接状态，离开页面后又显示未连接。图谱构建也只能逐个文档同步，无法稳定完成全量和增量构建。为了解决这个问题，系统需要把 Neo4j 健康状态提升为管理员共享状态，并让全量构建可以逐篇处理文档，单篇失败时记录原因而不是中断全部任务。",
    )
    add_paragraph(
        anchor,
        "第三类问题来自用户体验。侧边栏会随着主内容滚动而消失，管理员菜单功能重复，表格按钮换行，登录页和侧栏 Logo 与文字重叠，智能问答要等待完整回答后才显示结果。这些问题虽然不一定导致接口失败，却会影响实际使用。项目后期重点对导航、布局、问答流式输出、Markdown 渲染和报告同步进行了修复。",
    )

    add_paragraph(anchor, "2.2 概要设计", "h2")
    add_paragraph(
        anchor,
        "系统采用分层设计。浏览器端使用 Vue3 和 Element Plus 组织页面，负责登录注册、疾病预测、问答会话、图谱可视化、模型训练管理和审计日志展示。业务后端使用 Spring Boot 3，统一处理认证、权限、用户、知识文档、图谱任务、问答记录、报告生成和审计。模型服务使用 FastAPI，负责疾病风险预测、多模型训练、模型启用和重启加载。数据层由 MySQL 或本地 H2、Neo4j 和文件目录组成，分别保存结构化数据、医学知识图谱和报告模型文件。",
    )
    add_figure(anchor, figs["architecture"], "图2 MedRisk 系统总体架构")
    add_figure(anchor, figs["image2_architecture"], "图2 image2 绘制的云端总体架构图")
    add_paragraph(
        anchor,
        "这种方案的好处是边界清晰。前端只关心交互和展示，业务后端负责权限和数据一致性，模型服务独立承担训练和推理，Neo4j 专门服务医学实体关系检索。与把所有功能写在一个后端服务中的方案相比，分层方案更容易维护，也更符合云端部署时逐服务重建和检查的需要。",
    )
    add_figure(anchor, figs["modules"], "图3 MedRisk 功能模块划分")
    add_figure(anchor, figs["image2_modules"], "图3 image2 绘制的系统功能框图")

    add_paragraph(anchor, "2.3 详细设计与核心实现", "h2")
    add_paragraph(
        anchor,
        "疾病预测模块采用简洁输入和结构化输出。用户填写年龄、血糖、BMI、血压、吸烟、心脏病史等指标后，前端把数据提交到后端，后端调用模型服务得到多疾病风险概率。结果页展示风险等级、模型置信度、主要影响因素和报告生成入口。管理员也可以进入该页面，用于测试普通用户的预测流程。",
    )
    add_figure(anchor, figs["prediction_flow"], "图4 疾病风险预测业务流程")
    add_figure(anchor, figs["image2_prediction"], "图4 疾病预测 UML 活动图")
    add_figure(anchor, screenshot("04-prediction-form-80.png"), "图5 疾病预测表单页面")
    add_figure(anchor, screenshot("05-prediction-result-80.png"), "图6 疾病预测结果与解释页面")

    add_paragraph(
        anchor,
        "知识图谱模块围绕文档构建和 3D 可视化设计。构建时先清理 MedRisk 写入的旧图谱，再按文档逐篇抽取 Document、Section、Keyword、Disease、Symptom、RiskFactor、Exam、Treatment 和 Drug 等节点。每篇文档使用文档级唯一 key，避免不同文档因为同名实体被过度合并。构建失败时记录具体原因，如 Neo4j 连接失败、空内容、非法字段或写入失败。这样管理员可以看到失败原因，而不是只看到构建失败四个字。",
    )
    add_figure(anchor, figs["graph_flow"], "图7 知识图谱全量构建流程")
    add_figure(anchor, figs["image2_graph"], "图7 image2 绘制的知识图谱构建流程图")
    add_figure(anchor, screenshot("08-knowledge-graph-80.png"), "图8 医学知识图谱三维可视化")

    add_paragraph(
        anchor,
        "智能问答模块经历了多次调整。最初问答只能返回完整结果，用户需要等待较长时间。后来增加了流式输出，发送后先显示已接受和思考中，随后把推理过程和回答正文一点一点追加到页面。系统还支持日常聊天和医学问答的自动判断。日常聊天不检索知识库，医学问答会检索知识图谱、疾病档案和病历案例。对于图片输入，系统自动选择支持多模态的百炼模型，文本问题则优先选择成本更低的文本模型。",
    )
    add_figure(anchor, figs["qa_flow"], "图9 智能问答与多模态处理流程")
    add_figure(anchor, figs["image2_qa"], "图9 智能问答与多模态时序图")
    add_figure(anchor, screenshot("07-qa-answer-80.png"), "图10 智能问答流式输出页面")

    add_paragraph(
        anchor,
        "模型训练管理模块从单一 XGBoost 扩展为多模型体系。默认可用模型包括 Logistic Regression、RandomForest、ExtraTrees、HistGradientBoosting 和 XGBoost。高级模型如 LightGBM、CatBoost、TabPFN、TabICL 和 FT Transformer 作为可选依赖，不安装时显示不可用原因，不影响基础模型训练。不同模型使用不同超参数表单，训练完成后保存模型类型、实际超参数、评估集来源、指标和启用状态。模型服务重启后会加载启用版本，保证预测接口使用当前平台选择的模型。",
    )
    add_figure(anchor, figs["image2_training"], "图11 模型训练与评估思维导图")
    add_figure(anchor, screenshot("09-training-management-80.png"), "图11 模型训练管理页面")
    add_figure(anchor, figs["metrics"], "图12 模型评估指标展示样式")

    add_paragraph(
        anchor,
        "审计日志模块用于记录平台关键操作。系统会解析 X Forwarded For、X Real IP 和远端地址，并把请求 IP 写入审计日志。登录、预测、删除会话、问答和系统管理操作都应尽量保留访问来源。为了应对异常访问，系统还增加了 IP 黑名单能力。这个设计让平台不仅能展示功能，也能说明如何追踪用户行为和保护系统安全。",
    )
    add_figure(anchor, screenshot("10-audit-log-80.png"), "图13 审计日志中的请求 IP 记录")

    add_paragraph(anchor, "2.4 编码测试与结果展示", "h2")
    add_paragraph(
        anchor,
        "编码完成后，我从本地和远端两个环境进行验证。本地 Windows 启动固定使用 conda MedRisk 环境，后端使用 demo profile，模型服务使用 FastAPI，前端使用 Vite。远端 med_tencent 通过 Docker Compose 运行 MySQL、Neo4j、后端、前端和模型服务。2026 年 7 月 1 日晚重新实测时，公网 80 端口可以正常打开前端，/api/health 返回 UP，因此终版报告截图优先使用 81.71.65.98:80。",
    )
    add_figure(anchor, screenshot("01-login-page-80.png"), "图14 MedRisk 80 入口登录页面")
    add_figure(anchor, screenshot("02-login-failed-80.png"), "图15 登录失败异常提示")
    add_figure(anchor, screenshot("03-admin-home-80.png"), "图16 管理员控制台页面")
    add_paragraph(
        anchor,
        "测试过程不是只看页面能不能打开，而是把后端测试、前端构建、模型服务测试、Compose 配置和远端健康检查合在一起确认。后端使用 mvn test 检查业务逻辑，前端使用 npm run test 和 npm run build 检查组件和生产构建，模型服务使用 conda MedRisk Python 执行 pytest，远端使用 docker compose config 和健康检查接口确认部署配置。最终健康接口返回 UP，说明云端主要服务已经可以响应。",
    )
    add_table_before(
        doc,
        anchor,
        [
            ["测试项", "输入或操作", "实际结果", "结论"],
            ["公网入口", "访问 http://81.71.65.98:80", "HTTP 200，返回 MedRisk 前端页面。", "通过"],
            ["健康检查", "访问 /api/health", "返回 code 0，status 为 UP，service 为 medrisk backend。", "通过"],
            ["登录异常", "管理员账号输入错误密码", "页面出现 401 登录失败提示。", "通过"],
            ["管理员登录", "admin 和演示密码登录", "进入管理控制台，显示 4 个用户、12 条高风险记录、7 个训练任务和 5 个启用模型。", "通过"],
            ["疾病预测", "提交年龄、血糖、BMI、血压和病史指标", "页面返回高风险结果，最大疾病概率约为 74%，并展示 SHAP 影响因素。", "通过"],
            ["智能问答", "询问高血压和糖尿病风险管理", "回答按 Markdown 渲染，页面显示模型 qwen plus 和证据来源。", "通过"],
            ["知识图谱", "进入图谱可视化", "显示 240 个疾病节点、251 个关系节点和三维网络布局。", "通过"],
            ["模型训练", "查看训练任务", "平台显示多模型训练任务、评估集、进度、指标和操作入口。", "通过"],
            ["审计日志", "查看最近操作", "登录、医学问答、疾病预测等记录都显示请求 IP。", "通过"],
            ["多模态降级", "包含图片的问答请求", "系统应自动选择全模态模型；免费额度或模型不可用时显示降级提示而不阻塞会话。", "设计通过"],
            ["Neo4j 降级", "图谱不可用时访问问答", "系统应提示检索不可用并继续使用疾病档案和本地兜底回答。", "设计通过"],
        ],
    )
    add_figure(anchor, figs["test"], "图16 测试与部署验证结果")
    add_figure(anchor, figs["image2_test"], "图17 image2 绘制的测试覆盖矩阵图")
    add_figure(anchor, screenshot("11-health-check-80.png"), "图18 云端后端健康检查结果", width=4.6)

    add_paragraph(anchor, "2.5 结束语", "h2")
    add_paragraph(anchor, "2.5.1 项目工作总结", "h3")
    add_paragraph(
        anchor,
        "本次项目实践完成了 MedRisk 平台从启动修复到云端展示的较完整迭代。系统能够完成用户登录、邮箱验证码、疾病风险预测、报告生成、知识图谱构建、三维图谱可视化、智能问答、多模态输入、模型训练管理、审计日志和健康检查。更重要的是，项目解决了多个真实工程问题，包括 Flyway 导致的启动失败、Neo4j 状态不一致、图谱批量构建不稳定、问答输出等待过久、页面布局遮挡、审计 IP 缺失和远端部署入口差异。",
    )
    add_paragraph(anchor, "2.5.2 工程伦理审查与评价", "h3")
    add_paragraph(
        anchor,
        "MedRisk 面向医学健康场景，设计时必须考虑社会、健康、安全、法律、文化和隐私影响。系统在页面和报告中说明平台仅用于教学演示和健康风险提示，不能替代医生诊断。邮箱运行凭据、API 访问凭据和访问 token 不写入源码、报告和文档。智能问答对医学问题展示证据来源，并在证据不足时提示不能作为诊断依据。审计日志记录请求 IP，便于追踪异常行为。平台的积极意义是帮助用户理解风险因素和医学知识，局限是模型输出受训练数据、输入质量和外部模型稳定性影响，需要专业医生参与解释。",
    )
    add_paragraph(anchor, "2.5.3 团队协作工作总结", "h3")
    add_paragraph(
        anchor,
        "项目推进中，我主要承担问题定位、后端与模型服务联调、前端页面优化、云端部署验证和报告整理等工作。每次需求变化都会影响多个模块，例如智能问答从普通回答变为流式输出后，后端接口、前端渲染、审计日志和报告同步都要一起调整。这个过程让我认识到团队项目需要及时记录需求、明确接口边界、保护他人已有改动，并通过测试和截图证明修改确实有效。",
    )
    add_paragraph(anchor, "2.5.4 职业素养总结", "h3")
    add_paragraph(
        anchor,
        "通过实习，我对计算机专业学生在工程项目中的责任有了更直接的认识。面对启动失败、接口异常和部署问题时，需要先收集证据，再判断根因，不能凭感觉修改。面对医疗场景和用户数据时，需要保持谨慎，不能夸大系统能力，也不能泄露密钥和个人信息。面对文档和答辩时，需要把问题、方案、实现和验证讲清楚，让别人能够复现和检查。这些经历提升了我的工程意识、表达能力和独立解决问题的能力。",
    )

    add_paragraph(anchor, "2.6 参考文献", "h2")
    add_reference_list(anchor)

    add_page_break(anchor)
    add_paragraph(anchor, "三、实习体会", "h1")
    reflections = [
        "这次生产实习让我第一次比较完整地经历了一个综合系统从想法到上线展示的过程。以前做课程作业时，我更多关注某一个功能是否可以运行，例如一个模型能不能训练，一个页面能不能显示，一个接口能不能返回数据。MedRisk 项目让我意识到，真实工程不只是把单个功能写出来，还要让许多功能在同一套环境里稳定协作。前端、后端、模型服务、数据库、知识图谱、云服务器和文档报告都会互相影响。只要其中一个环节没有处理好，用户看到的可能就是登录失败、页面空白或接口超时。",
        "实习中最有收获的是问题定位能力。项目启动失败时，浏览器只提示请求返回 500，如果只盯着前端就很难找到答案。继续查看后端日志后，才发现是本地 H2 数据库中的旧 Flyway 迁移记录和新的迁移文件不一致，导致后端没有真正启动。后来修复方案也不是简单删除数据库，而是把本地 demo profile 和 Linux 云端 MySQL 区分开来，让本地演示能够容忍历史 checksum 差异，同时保持云端迁移策略严格。这个过程让我理解了工程修复应该尽量缩小影响范围，不能为了快速启动而破坏生产环境规则。",
        "知识图谱和智能问答模块让我把理论学习与项目实践联系起来。第六天理论练习中学习的 Mini RAG，本质上就是先检索证据再生成回答。项目中的医学问答也采用类似思路，只是检索来源从课堂中的小型文档列表变成了 Neo4j 知识图谱、疾病档案和病历案例。第七天练习中的 GCN 和图结构可视化，让我更容易理解医学实体之间的关系为什么适合用图数据库保存。通过这种对照，我发现课堂练习并不是孤立内容，它们可以成为真实项目中设计方案的依据。",
        "界面优化也改变了我对前端工作的看法。最初我以为前端主要是把功能按钮放到页面上，但实际使用时，侧边栏滚动消失、按钮换行、Logo 遮挡文字、图谱宽度超过屏幕、问答输出等待过久，都会明显影响用户体验。后来我逐步调整侧边栏固定与收起、管理员菜单去重、按钮一行显示、问答流式输出和 Markdown 渲染。这个过程说明，界面不仅要好看，还要让用户知道自己在哪里、可以做什么、操作后发生了什么。",
        "医疗健康场景让我更加重视工程伦理。疾病风险预测和智能问答可能会影响用户对自身健康的判断，所以系统必须清楚说明它只能用于教学演示和健康风险提示，不能替代医生诊断。大模型回答也需要给出证据来源和边界提醒。另一方面，系统接入邮箱验证码、阿里云百炼 API 和远端服务器时，会涉及密钥、访问凭据和访问 token，这些内容不能写入源码和报告。实习过程中我反复检查文档和交付目录，避免泄露真实密钥。这让我认识到安全意识不是最后才补的工作，而应该贯穿设计、编码、测试和文档全过程。",
        "最后两天整理报告和答辩材料时，我又重新梳理了整个项目。写报告并不是简单记录做过什么，而是把问题背景、设计思路、实现过程、测试结果和个人体会连接起来。通过整理截图、流程图、架构图和测试记录，我能够更清楚地看到自己的工作轨迹，也能发现哪些地方还可以继续完善。后续如果继续迭代 MedRisk，我希望进一步优化多疾病统一预测模型，补充更多公开医学数据集评估，完善 IP 黑名单和端到端自动化测试，让平台更加稳定、可信和易维护。",
    ]
    for item in reflections:
        add_paragraph(anchor, item)


def scrub_check(text: str) -> list[str]:
    patterns = [
        r"sk-[A-Za-z0-9_.-]{12,}",
        r"access_token",
        r"refresh_token",
        r"Bearer\s+[A-Za-z0-9_.-]+",
        r"授权码\s*[A-Za-z0-9_.-]+",
        r"password\s*[:=]\s*[^\\s]+",
    ]
    hits: list[str] = []
    for pattern in patterns:
        if re.search(pattern, text, flags=re.IGNORECASE):
            hits.append(pattern)
    return hits


def main() -> None:
    if not BASE_DOCX.exists():
        candidates = [p for p in FINAL_DIR.glob("*.docx") if not p.name.startswith("~$")]
        if not candidates:
            raise FileNotFoundError("No base DOCX found in final_ret")
        base = candidates[0]
    else:
        base = BASE_DOCX

    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(str(base))
    anchor_element = remove_blocks_between(doc, "一、实习任务与计划", "I职业素养评分表")
    anchor = paragraph_from_element(doc, anchor_element)
    add_report_content(doc, anchor)

    doc.save(str(OUT_DOCX))

    # Keep a render input copy for QA and make notes explicit.
    shutil.copy2(OUT_DOCX, RENDER_DIR / "render_input_image2_final.docx")
    all_text = "\n".join(p.text for p in Document(str(OUT_DOCX)).paragraphs)
    secrets = scrub_check(all_text)
    summary = (
        "# image2 图文增强终版生产实习报告生成记录\n\n"
        f"- 基底 DOCX：`{base}`\n"
        f"- 输出 DOCX：`{OUT_DOCX}`\n"
        "- 时间线：2026.6.15 实习动员，2026.6.16 到 2026.6.22 理论学习，2026.6.23 到 2026.7.3 项目开发与管理，2026.7.4 到 2026.7.5 报告与答辩。\n"
        "- 截图入口：公网 80 端口本轮实测返回 200，健康检查接口返回 UP，终版截图优先使用 81.71.65.98:80。\n"
        "- 图文材料：包含 image2 生成并后处理的功能框图、总体架构图、项目流程图、知识图谱流程图、问答时序图、疾病预测 UML、训练评估思维导图、测试覆盖矩阵，以及远端系统截图。\n"
        f"- 保密扫描：{'发现疑似敏感模式 ' + ', '.join(secrets) if secrets else '未发现真实密钥、访问凭据、token 或明文密码模式'}。\n"
    )
    (NOTES_DIR / "image2-final-generation-summary.md").write_text(summary, encoding="utf-8")
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
